from flask import Flask, render_template, request, jsonify, send_file
import xml.etree.ElementTree as ET
import os
import json
import csv
import io
import tempfile
from werkzeug.utils import secure_filename
import signal
import sys
from datetime import datetime

# Importaciones para exportación de diferentes formatos
try:
    import openpyxl
    from openpyxl.styles import Font, Alignment, PatternFill
    EXCEL_AVAILABLE = True
except ImportError:
    EXCEL_AVAILABLE = False

try:
    from docx import Document
    from docx.shared import Inches
    DOC_AVAILABLE = True
except ImportError:
    DOC_AVAILABLE = False

try:
    from reportlab.lib.pagesizes import letter, A4
    from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib import colors
    from reportlab.lib.units import inch
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False

# Determinar las rutas correctas para templates y static
import os
current_dir = os.path.dirname(os.path.abspath(__file__))

# Configurar las rutas de templates y static
template_dir = os.path.join(current_dir, 'templates')
static_dir = os.path.join(current_dir, 'static')
upload_dir = os.path.join(current_dir, 'uploads')
translations_dir = os.path.join(current_dir, 'translations')

# Verificar que existan los directorios, si no, crearlos
os.makedirs(template_dir, exist_ok=True)
os.makedirs(static_dir, exist_ok=True)
os.makedirs(upload_dir, exist_ok=True)
os.makedirs(translations_dir, exist_ok=True)

print(f"Directorio actual: {current_dir}")
print(f"Directorio de templates: {template_dir}")
print(f"Directorio de static: {static_dir}")
print(f"Directorio de uploads: {upload_dir}")
print(f"Directorio de traducciones: {translations_dir}")

app = Flask(__name__, 
            template_folder=template_dir,
            static_folder=static_dir)
app.config['MAX_CONTENT_LENGTH'] = 16 * 1024 * 1024  # 16MB max file size
app.config['UPLOAD_FOLDER'] = upload_dir

# Global variables for translations
translations_cache = {}
language_config = {}

def load_translations():
    """Load all translation files and language configuration"""
    global translations_cache, language_config
    
    # Load language configuration
    config_path = os.path.join(translations_dir, 'config.json')
    try:
        with open(config_path, 'r', encoding='utf-8') as f:
            language_config = json.load(f)
    except FileNotFoundError:
        print(f"Language config not found at {config_path}")
        language_config = {"languages": [{"code": "es", "name": "Español"}], "defaultLanguage": "es"}
    
    # Load all translation files
    for lang_info in language_config.get('languages', []):
        lang_code = lang_info['code']
        lang_path = os.path.join(translations_dir, f'{lang_code}.json')
        try:
            with open(lang_path, 'r', encoding='utf-8') as f:
                translations_cache[lang_code] = json.load(f)
        except FileNotFoundError:
            print(f"Translation file not found for {lang_code} at {lang_path}")

def get_translation(key, lang='es', fallback=None, **kwargs):
    """Get a translation for a given key and language"""
    if lang not in translations_cache:
        lang = language_config.get('fallbackLanguage', 'es')
    
    translation = translations_cache.get(lang, {})
    
    # Handle nested keys (e.g., 'table_headers.device')
    keys = key.split('.')
    for k in keys:
        if isinstance(translation, dict) and k in translation:
            translation = translation[k]
        else:
            # Fallback to English or Spanish
            fallback_lang = 'en' if lang != 'en' else 'es'
            if fallback_lang in translations_cache:
                fallback_translation = translations_cache[fallback_lang]
                for k in keys:
                    if isinstance(fallback_translation, dict) and k in fallback_translation:
                        fallback_translation = fallback_translation[k]
                    else:
                        return fallback or key  # Return fallback or key if no translation found
                translation = fallback_translation
            else:
                return fallback or key
            break
    
    # Replace placeholders if kwargs provided
    if isinstance(translation, str) and kwargs:
        try:
            # Handle both {placeholder} and {language} style placeholders
            for key, value in kwargs.items():
                translation = translation.replace(f'{{{key}}}', str(value))
        except Exception as e:
            print(f"Error replacing placeholders in translation '{translation}': {e}")
            pass  # Ignore missing placeholders
    
    return translation

def parse_shw_file(filepath):
    """Parses the .shw file and extracts device information."""
    devices = []
    try:
        tree = ET.parse(filepath)
        root = tree.getroot()
        
        # Find all device elements
        all_devices = root.findall('.//device')
        
        for i, device in enumerate(all_devices):
            device_info = {}
            
            # Extract top-level device details
            series = device.find('series')
            if series is not None and series.text:
                device_info['series'] = series.text.strip()
            else:
                device_info['series'] = 'N/A'

            model = device.find('model')
            if model is not None and model.text:
                device_info['model'] = model.text.strip()
            else:
                device_info['model'] = 'N/A'

            device_name_elem = device.find('device_name')
            if device_name_elem is not None and device_name_elem.text:
                 device_info['device_name'] = device_name_elem.text.strip()
            else:
                 device_info['device_name'] = 'N/A'

            zone_elem = device.find('zone')
            if zone_elem is not None and zone_elem.text:
                device_info['zone'] = zone_elem.text.strip()
            else:
                device_info['zone'] = 'N/A'

            band_elem = device.find('band')
            if band_elem is not None and band_elem.text:
                device_info['band'] = band_elem.text.strip()
            else:
                device_info['band'] = 'N/A'

            device_info['channels'] = []

            # Find all channel elements within the device
            channels = device.findall('channel')
            
            for channel in channels:
                channel_info = {}
                
                channel_name_elem = channel.find('channel_name')
                if channel_name_elem is not None:
                    # Handle CDATA sections
                    if channel_name_elem.text:
                        channel_info['name'] = channel_name_elem.text.strip()
                    else:
                        channel_info['name'] = 'N/A'
                else:
                    channel_info['name'] = 'N/A'

                frequency_elem = channel.find('frequency')
                if frequency_elem is not None and frequency_elem.text:
                    try:
                        # Frequency is in kHz, convert to MHz for display
                        freq_mhz = int(frequency_elem.text) / 1000
                        channel_info['frequency'] = f"{freq_mhz:.3f}"
                    except (ValueError, TypeError):
                        channel_info['frequency'] = 'N/A'
                else:
                    channel_info['frequency'] = 'N/A'
                
                # Always add channel, even if some info is missing
                device_info['channels'].append(channel_info)

            # Always add device if it has any meaningful data
            if (device_info.get('device_name') != 'N/A' or 
                device_info.get('model') != 'N/A' or 
                device_info.get('series') != 'N/A' or 
                device_info['channels']):
                devices.append(device_info)

        return devices

    except ET.ParseError as e:
        print(f"Error parsing XML: {e}")
        return []
    except FileNotFoundError:
        print(f"File not found: {filepath}")
        return []
        
    return devices

@app.route('/')
def index():
    return render_template('index.html')

@app.route('/api/languages')
def get_languages():
    """Get available languages"""
    return jsonify({
        'languages': language_config.get('languages', []),
        'defaultLanguage': language_config.get('defaultLanguage', 'es')
    })

@app.route('/api/translations/<lang>')
def get_translations(lang):
    """Get translations for a specific language"""
    if lang in translations_cache:
        return jsonify(translations_cache[lang])
    else:
        # Return fallback language
        fallback_lang = language_config.get('fallbackLanguage', 'es')
        return jsonify(translations_cache.get(fallback_lang, {}))

@app.route('/api/set_language', methods=['POST'])
def set_language():
    """Set user's preferred language (for future server-side features)"""
    data = request.get_json()
    lang = data.get('language', 'es')
    
    # Get the native name of the language
    language_names = {
        'es': 'Español',
        'en': 'English',
        'fr': 'Français',
        'de': 'Deutsch',
        'it': 'Italiano',
        'pt': 'Português',
        'ca': 'Català',
        'gl': 'Galego',
        'eu': 'Euskera'
    }
    language_name = language_names.get(lang, lang)
    
    # In a real app, you might save this to user preferences
    # For now, just acknowledge the change
    return jsonify({
        'success': True,
        'language': lang,
        'message': get_translation('messages.language_changed', lang, language=language_name)
    })

@app.route('/load_file_path', methods=['POST'])
def load_file_path():
    """Load a .shw file from a given file path (for Electron file dialog)"""
    data = request.get_json()
    file_path = data.get('file_path')
    
    if not file_path:
        return jsonify({'error': 'No file path provided'}), 400
    
    if not os.path.exists(file_path):
        return jsonify({'error': 'File not found'}), 404
    
    if not file_path.lower().endswith('.shw'):
        return jsonify({'error': 'Invalid file type. Only .shw files are allowed.'}), 400
    
    try:
        # Parse the file directly from the path
        devices = parse_shw_file(file_path)
        
        # Convert to flat structure for table display
        rows = []
        for device in devices:
            device_display_name = f"{device.get('device_name', 'Dispositivo')} ({device.get('model', 'N/A')})"
            zone = device.get('zone', 'N/A')
            band = device.get('band', 'N/A')

            if device['channels']:
                for i, channel in enumerate(device['channels']):
                    name = channel.get('name', 'N/A')
                    freq = channel.get('frequency', 'N/A')
                    
                    # For the first channel, show device name. For others, leave it blank.
                    if i == 0:
                        rows.append({
                            'id': f"{len(rows)}",
                            'device': device_display_name,
                            'channel': name,
                            'frequency': freq,
                            'zone': zone,
                            'band': band
                        })
                    else:
                        rows.append({
                            'id': f"{len(rows)}",
                            'device': '',
                            'channel': name,
                            'frequency': freq,
                            'zone': zone,
                            'band': band
                        })
            else:
                # If a device has no channels, still show the device
                rows.append({
                    'id': f"{len(rows)}",
                    'device': device_display_name,
                    'channel': 'N/A',
                    'frequency': 'N/A',
                    'zone': zone,
                    'band': band
                })
        
        return jsonify({
            'success': True,
            'devices': len(devices),
            'rows': len(rows),
            'data': rows,
            'filename': os.path.basename(file_path)
        })
        
    except Exception as e:
        return jsonify({'error': f'Error processing file: {str(e)}'}), 500

@app.route('/upload', methods=['POST'])
def upload_file():
    if 'file' not in request.files:
        return jsonify({'error': 'No file selected'}), 400
    
    file = request.files['file']
    
    if file.filename == '':
        return jsonify({'error': 'No file selected'}), 400
    
    if file and file.filename and file.filename.lower().endswith('.shw'):
        # Usar directorio temporal para archivos
        with tempfile.NamedTemporaryFile(delete=False, suffix='.shw') as temp_file:
            file.save(temp_file.name)
            filepath = temp_file.name
        
        try:
            # Parse the file
            devices = parse_shw_file(filepath)
            
            # Convert to flat structure for table display
            rows = []
            for device in devices:
                device_display_name = f"{device.get('device_name', 'Dispositivo')} ({device.get('model', 'N/A')})"
                zone = device.get('zone', 'N/A')
                band = device.get('band', 'N/A')

                if device['channels']:
                    for i, channel in enumerate(device['channels']):
                        name = channel.get('name', 'N/A')
                        freq = channel.get('frequency', 'N/A')
                        
                        # For the first channel, show device name. For others, leave it blank.
                        if i == 0:
                            rows.append({
                                'id': f"{len(rows)}",
                                'device': device_display_name,
                                'channel': name,
                                'frequency': freq,
                                'zone': zone,
                                'band': band
                            })
                        else:
                            rows.append({
                                'id': f"{len(rows)}",
                                'device': '',
                                'channel': name,
                                'frequency': freq,
                                'zone': zone,
                                'band': band
                            })
                else:
                    # If a device has no channels, still show the device
                    rows.append({
                        'id': f"{len(rows)}",
                        'device': device_display_name,
                        'channel': 'N/A',
                        'frequency': 'N/A',
                        'zone': zone,
                        'band': band
                    })
            
            return jsonify({
                'success': True,
                'devices': len(devices),
                'rows': len(rows),
                'data': rows
            })
            
        finally:
            # Clean up temporary file
            try:
                os.unlink(filepath)
            except:
                pass
    
    return jsonify({'error': 'Invalid file type. Only .shw files are allowed.'}), 400

@app.route('/update_row', methods=['POST'])
def update_row():
    """Update a row's data"""
    data = request.get_json()
    # In a real application, you would save this to a database
    # For now, just return success
    return jsonify({'success': True})

@app.route('/export_csv', methods=['POST'])
def export_csv():
    """Export table data to CSV"""
    data = request.get_json()
    rows = data.get('rows', [])
    lang = data.get('language', 'es')
    
    if not rows:
        error_msg = get_translation('messages.no_data', lang)
        return jsonify({'error': error_msg}), 400
    
    # Create CSV content
    output = io.StringIO()
    writer = csv.writer(output)
    
    # Write header with translations
    headers = [
        get_translation('table_headers.channel', lang),
        get_translation('table_headers.frequency', lang),
        get_translation('table_headers.rf_zone', lang),
        get_translation('table_headers.device_model', lang),
        get_translation('table_headers.band', lang)
    ]
    writer.writerow(headers)
    
    # Write data
    for row in rows:
        freq_display = f"{row['frequency']} MHz" if row['frequency'] != 'N/A' else 'N/A'
        writer.writerow([
            row['channel'], 
            freq_display,
            row['zone'],
            row['device'],
            row['band']
        ])
    
    csv_content = output.getvalue()
    output.close()
    
    return jsonify({
        'success': True,
        'csv_data': csv_content,
        'filename': 'shw_data_export.csv'
    })

@app.route('/export_excel', methods=['POST'])
def export_excel():
    """Export table data to Excel"""
    if not EXCEL_AVAILABLE:
        return jsonify({'error': 'Excel export not available. Install openpyxl.'}), 500
    
    data = request.get_json()
    rows = data.get('rows', [])
    lang = data.get('language', 'es')
    
    if not rows:
        error_msg = get_translation('messages.no_data', lang)
        return jsonify({'error': error_msg}), 400
    
    # Create workbook
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "SHW Data"
    
    # Headers with translations
    headers = [
        get_translation('table_headers.channel', lang),
        get_translation('table_headers.frequency', lang),
        get_translation('table_headers.rf_zone', lang),
        get_translation('table_headers.device_model', lang),
        get_translation('table_headers.band', lang)
    ]
    
    # Write headers with styling
    for col, header in enumerate(headers, 1):
        cell = ws.cell(row=1, column=col, value=header)
        cell.font = Font(bold=True, color="FFFFFF")
        cell.fill = PatternFill(start_color="366092", end_color="366092", fill_type="solid")
        cell.alignment = Alignment(horizontal="center", vertical="center")
    
    # Write data
    for row_idx, row in enumerate(rows, 2):
        freq_display = f"{row['frequency']} MHz" if row['frequency'] != 'N/A' else 'N/A'
        data_row = [
            row['channel'],
            freq_display,
            row['zone'],
            row['device'],
            row['band']
        ]
        
        for col, value in enumerate(data_row, 1):
            cell = ws.cell(row=row_idx, column=col, value=value)
            cell.alignment = Alignment(horizontal="left", vertical="center")
    
    # Auto-adjust column widths
    for column in ws.columns:
        max_length = 0
        column_letter = column[0].column_letter
        for cell in column:
            try:
                if len(str(cell.value)) > max_length:
                    max_length = len(str(cell.value))
            except:
                pass
        adjusted_width = min(max_length + 2, 50)
        ws.column_dimensions[column_letter].width = adjusted_width
    
    # Save to temporary file
    temp_file = tempfile.NamedTemporaryFile(delete=False, suffix='.xlsx')
    wb.save(temp_file.name)
    temp_file.close()
    
    try:
        return send_file(
            temp_file.name,
            as_attachment=True,
            download_name=f'shw_data_export_{datetime.now().strftime("%Y%m%d_%H%M%S")}.xlsx',
            mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet'
        )
    finally:
        # Clean up temp file after a delay
        def cleanup():
            try:
                os.unlink(temp_file.name)
            except:
                pass
        
        # Schedule cleanup for later
        import threading
        timer = threading.Timer(10.0, cleanup)
        timer.start()

@app.route('/export_doc', methods=['POST'])
def export_doc():
    """Export table data to Word document"""
    if not DOC_AVAILABLE:
        return jsonify({'error': 'Word export not available. Install python-docx.'}), 500
    
    data = request.get_json()
    rows = data.get('rows', [])
    lang = data.get('language', 'es')
    
    if not rows:
        error_msg = get_translation('messages.no_data', lang)
        return jsonify({'error': error_msg}), 400
    
    # Create document
    doc = Document()
    
    # Add title
    title = doc.add_heading('SHW Reader - ' + get_translation('export.data_export', lang, fallback='Data Export'), 0)
    
    # Add export info
    info_para = doc.add_paragraph()
    info_para.add_run(f'{get_translation("export.generated_on", lang, fallback="Generated on")}: ')
    info_para.add_run(datetime.now().strftime('%Y-%m-%d %H:%M:%S')).bold = True
    info_para.add_run(f'\n{get_translation("export.total_records", lang, fallback="Total records")}: ');
    info_para.add_run(str(len(rows))).bold = True
    
    # Add some space
    doc.add_paragraph()
    
    # Headers with translations
    headers = [
        get_translation('table_headers.channel', lang),
        get_translation('table_headers.frequency', lang),
        get_translation('table_headers.rf_zone', lang),
        get_translation('table_headers.device_model', lang),
        get_translation('table_headers.band', lang)
    ]
    
    # Create table
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'
    
    # Add header row
    hdr_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        hdr_cells[i].text = header
        # Make header bold
        for paragraph in hdr_cells[i].paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
    
    # Add data rows
    for row in rows:
        freq_display = f"{row['frequency']} MHz" if row['frequency'] != 'N/A' else 'N/A'
        data_row = [
            row['channel'],
            freq_display,
            row['zone'],
            row['device'],
            row['band']
        ]
        
        row_cells = table.add_row().cells
        for i, value in enumerate(data_row):
            row_cells[i].text = str(value)
    
    # Set column widths
    for column in table.columns:
        for cell in column.cells:
            cell.width = Inches(1.5)
    
    # Save to temporary file
    temp_file = tempfile.NamedTemporaryFile(delete=False, suffix='.docx')
    doc.save(temp_file.name)
    temp_file.close()
    
    try:
        return send_file(
            temp_file.name,
            as_attachment=True,
            download_name=f'shw_data_export_{datetime.now().strftime("%Y%m%d_%H%M%S")}.docx',
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
    finally:
        # Clean up temp file after a delay
        def cleanup():
            try:
                os.unlink(temp_file.name)
            except:
                pass
        
        import threading
        timer = threading.Timer(10.0, cleanup)
        timer.start()

@app.route('/export_pdf', methods=['POST'])
def export_pdf():
    """Export table data to PDF"""
    if not PDF_AVAILABLE:
        return jsonify({'error': 'PDF export not available. Install reportlab.'}), 500
    
    data = request.get_json()
    rows = data.get('rows', [])
    lang = data.get('language', 'es')
    
    if not rows:
        error_msg = get_translation('messages.no_data', lang)
        return jsonify({'error': error_msg}), 400
    
    # Create temporary file
    temp_file = tempfile.NamedTemporaryFile(delete=False, suffix='.pdf')
    temp_file.close()
    
    # Create PDF document
    doc = SimpleDocTemplate(temp_file.name, pagesize=A4)
    elements = []
    
    # Get styles
    styles = getSampleStyleSheet()
    title_style = ParagraphStyle(
        'CustomTitle',
        parent=styles['Heading1'],
        fontSize=18,
        textColor=colors.darkblue,
        alignment=1  # Center alignment
    )
    
    # Add title
    title = Paragraph('SHW Reader - ' + get_translation('export.data_export', lang, fallback='Data Export'), title_style)
    elements.append(title)
    elements.append(Spacer(1, 20))
    
    # Add export info
    info_text = f"""
    <b>{get_translation('export.generated_on', lang, fallback='Generated on')}:</b> {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}<br/>
    <b>{get_translation('export.total_records', lang, fallback='Total records')}:</b> {len(rows)}
    """
    info = Paragraph(info_text, styles['Normal'])
    elements.append(info)
    elements.append(Spacer(1, 20))
    
    # Headers with translations
    headers = [
        get_translation('table_headers.channel', lang),
        get_translation('table_headers.frequency', lang),
        get_translation('table_headers.rf_zone', lang),
        get_translation('table_headers.device_model', lang),
        get_translation('table_headers.band', lang)
    ]
    
    # Prepare table data
    table_data = [headers]
    
    for row in rows:
        freq_display = f"{row['frequency']} MHz" if row['frequency'] != 'N/A' else 'N/A'
        table_data.append([
            row['channel'],
            freq_display,
            row['zone'],
            row['device'],
            row['band']
        ])
    
    # Create table
    table = Table(table_data)
    table.setStyle(TableStyle([
        # Header styling
        ('BACKGROUND', (0, 0), (-1, 0), colors.darkblue),
        ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
        ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
        ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
        ('FONTSIZE', (0, 0), (-1, 0), 10),
        ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
        
        # Data styling
        ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
        ('TEXTCOLOR', (0, 1), (-1, -1), colors.black),
        ('FONTNAME', (0, 1), (-1, -1), 'Helvetica'),
        ('FONTSIZE', (0, 1), (-1, -1), 8),
        ('ALIGN', (0, 1), (-1, -1), 'LEFT'),
        
        # Grid
        ('GRID', (0, 0), (-1, -1), 1, colors.black),
        ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
    ]))
    
    elements.append(table)
    
    # Build PDF
    doc.build(elements)
    
    try:
        return send_file(
            temp_file.name,
            as_attachment=True,
            download_name=f'shw_data_export_{datetime.now().strftime("%Y%m%d_%H%M%S")}.pdf',
            mimetype='application/pdf'
        )
    finally:
        # Clean up temp file after a delay
        def cleanup():
            try:
                os.unlink(temp_file.name)
            except:
                pass
        
        import threading
        timer = threading.Timer(10.0, cleanup)
        timer.start()

def signal_handler(sig, frame):
    print('\\nCerrando servidor Flask...')
    sys.exit(0)

if __name__ == '__main__':
    # Configurar manejador de señales para cierre limpio
    signal.signal(signal.SIGINT, signal_handler)
    signal.signal(signal.SIGTERM, signal_handler)
    
    # Cargar traducciones antes de iniciar el servidor
    print("🌐 Cargando traducciones...")
    load_translations()
    print(f"✅ Traducciones cargadas para {len(translations_cache)} idiomas")
    
    port = int(os.environ.get('PORT', 5001))
    print("🚀 Iniciando SHW Reader Server...")
    print(f"🌐 Servidor disponible en: http://localhost:{port}")
    
    # Ejecutar en modo de producción para la app de escritorio
    app.run(host='127.0.0.1', port=port, debug=False, use_reloader=False)
